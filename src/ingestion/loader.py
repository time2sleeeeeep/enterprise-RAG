# 文档加载模块：支持 PDF（fitz）、Markdown、DOCX 三种格式的文档加载。
# PDF 页面添加 [PAGE N] 标记以便后续按页分块；统一返回 LoadedDocument 数据类。

from pathlib import Path
from dataclasses import dataclass, field

import fitz
from docx import Document as DocxDocument


@dataclass
class LoadedDocument:
    content: str
    filename: str
    file_type: str
    metadata: dict = field(default_factory=dict)


def load_pdf(file_path: Path) -> LoadedDocument:
    """使用 PyMuPDF 逐页提取文本，非空页加 [PAGE N] 标记后拼接为整体内容。"""
    doc = fitz.open(str(file_path))
    pages = []
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text("text")
        if text.strip():
            pages.append(f"[PAGE {page_num}]\n{text}")
    doc.close()
    return LoadedDocument(
        content="\n\n".join(pages),
        filename=file_path.name,
        file_type="pdf",
        metadata={"page_count": len(pages)},
    )


def load_markdown(file_path: Path) -> LoadedDocument:
    """以 UTF-8 读取 Markdown 文件全文，无分页处理。"""
    content = file_path.read_text(encoding="utf-8")
    return LoadedDocument(
        content=content,
        filename=file_path.name,
        file_type="markdown",
    )


def load_docx(file_path: Path) -> LoadedDocument:
    """使用 python-docx 提取非空段落，段落间以双换行分隔。"""
    doc = DocxDocument(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return LoadedDocument(
        content="\n\n".join(paragraphs),
        filename=file_path.name,
        file_type="docx",
    )


LOADERS = {
    ".pdf": load_pdf,
    ".md": load_markdown,
    ".markdown": load_markdown,
    ".docx": load_docx,
}


def load_document(file_path: Path) -> LoadedDocument:
    """根据文件扩展名分发到对应加载器，不支持的格式抛出 ValueError。"""
    suffix = file_path.suffix.lower()
    loader = LOADERS.get(suffix)
    if loader is None:
        raise ValueError(f"Unsupported file type: {suffix}")
    return loader(file_path)
